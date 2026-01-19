"""
Word document generation for sprint reports
"""

from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger(__name__)


class WordDocumentGenerator:
    """Generate Word documents for sprint reports using templates"""

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize Word generator

        Args:
            template_path: Path to .docx template file (optional)
        """
        self.template_path = template_path
        self.document = None

    def load_template(self) -> Document:
        """Load template or create blank document"""
        if self.template_path and self.template_path.exists():
            logger.info(f"Loading template: {self.template_path}")
            try:
                return Document(str(self.template_path))
            except Exception as e:
                logger.warning(f"Failed to load template: {e}")
                logger.info("Falling back to blank document")
                return Document()
        else:
            logger.info("No template specified, using blank document")
            return Document()

    def generate(self, sprint_data: Dict[str, Any], metrics: Any,
                 ai_insights: Optional[Dict] = None) -> Document:
        """
        Generate complete Word document from sprint data

        Args:
            sprint_data: Sprint configuration (sprint_number, board_name, dates)
            metrics: Calculated sprint metrics
            ai_insights: Optional AI-generated insights

        Returns:
            Complete Word document
        """
        self.document = self.load_template()

        # Build document sections
        self._add_header_section(sprint_data)
        self._add_executive_summary(metrics, ai_insights)
        self._add_metrics_section(metrics)
        self._add_team_contributions(metrics)
        self._add_work_streams(metrics)

        if ai_insights:
            self._add_ai_insights_section(ai_insights)

        self._add_footer_section(sprint_data)

        return self.document

    def _add_header_section(self, sprint_data: Dict[str, Any]) -> None:
        """Add report header with sprint info"""
        # Title - use add_paragraph instead of add_heading to avoid style dependencies
        title = self.document.add_paragraph(f"Sprint {sprint_data['sprint_number']} Report")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Make title bold and large
        run = title.runs[0]
        run.bold = True
        run.font.size = Pt(20)

        # Metadata table - don't set style to avoid dependency on template styles
        table = self.document.add_table(rows=4, cols=2)
        table.style = 'Table Grid'  # Use basic built-in style

        cells = [
            ('Board:', sprint_data['board_name']),
            ('Sprint:', str(sprint_data['sprint_number'])),
            ('Start Date:', sprint_data.get('start_date', 'N/A')),
            ('End Date:', sprint_data.get('end_date', 'N/A'))
        ]

        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        self.document.add_paragraph()  # Spacing

    def _add_executive_summary(self, metrics: Any,
                               ai_insights: Optional[Dict] = None) -> None:
        """Add executive summary section"""
        self.document.add_heading('Executive Summary', level=1)

        if ai_insights and ai_insights.get('executive_summary'):
            # Use AI-generated summary
            self.document.add_paragraph(ai_insights['executive_summary'])
        else:
            # Fallback to rule-based summary
            summary = (
                f"Sprint completed with {metrics.done_count} of "
                f"{metrics.total_issues} issues finished "
                f"({metrics.completion_rate:.1f}% completion rate)."
            )
            self.document.add_paragraph(summary)

        self.document.add_paragraph()

    def _add_metrics_section(self, metrics: Any) -> None:
        """Add key metrics section with formatted table"""
        self.document.add_heading('Key Metrics', level=1)

        # Metrics table
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        metrics_data = [
            ('Total Issues', str(metrics.total_issues)),
            ('Completed', f"{metrics.done_count} ({metrics.completion_rate:.1f}%)"),
            ('In Progress', str(metrics.in_progress_count)),
            ('Not Started', str(metrics.not_started_count)),
            ('Story Points Completed', f"{metrics.completed_points}/{metrics.total_points}")
        ]

        for i, (label, value) in enumerate(metrics_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold the label
            row.cells[0].paragraphs[0].runs[0].bold = True

        self.document.add_paragraph()

    def _add_team_contributions(self, metrics: Any) -> None:
        """Add team contributions section"""
        self.document.add_heading('Team Contributions', level=1)

        if not hasattr(metrics, 'assignee_stats') or not metrics.assignee_stats:
            self.document.add_paragraph('No assignee data available.')
            return

        # Sort by completion count
        sorted_assignees = sorted(
            metrics.assignee_stats.items(),
            key=lambda x: x[1].get('done_count', 0),
            reverse=True
        )

        # Create table
        table = self.document.add_table(
            rows=len(sorted_assignees) + 1,
            cols=4
        )
        table.style = 'Light Grid Accent 1'

        # Header row
        headers = ['Team Member', 'Completed', 'In Progress', 'Total']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        for i, (assignee, stats) in enumerate(sorted_assignees, start=1):
            row = table.rows[i]
            row.cells[0].text = assignee
            row.cells[1].text = str(stats.get('done_count', 0))
            row.cells[2].text = str(stats.get('in_progress_count', 0))
            row.cells[3].text = str(stats.get('total', 0))

        self.document.add_paragraph()

    def _add_work_streams(self, metrics: Any) -> None:
        """Add work streams / epic progress section"""
        self.document.add_heading('Work Streams (Epics)', level=1)

        if not hasattr(metrics, 'epic_progress') or not metrics.epic_progress:
            self.document.add_paragraph('No epic data available.')
            return

        for epic_name, stats in metrics.epic_progress.items():
            # Epic subheading
            self.document.add_heading(epic_name or 'No Epic', level=2)

            # Epic stats
            completion = (stats.get('done_count', 0) / stats.get('total', 1)) * 100
            p = self.document.add_paragraph()
            p.add_run(f"Completion: ").bold = True
            p.add_run(f"{stats.get('done_count', 0)}/{stats.get('total', 0)} ")
            p.add_run(f"({completion:.1f}%)")

            p = self.document.add_paragraph()
            p.add_run(f"Story Points: ").bold = True
            p.add_run(f"{stats.get('completed_points', 0)}/{stats.get('total_points', 0)}")

        self.document.add_paragraph()

    def _add_ai_insights_section(self, ai_insights: Dict) -> None:
        """Add AI-generated insights section"""
        self.document.add_heading('AI Insights', level=1)

        if ai_insights.get('key_decisions'):
            self.document.add_heading('Key Decisions', level=2)
            self.document.add_paragraph(ai_insights['key_decisions'])

        if ai_insights.get('meeting_themes'):
            self.document.add_heading('Meeting Themes', level=2)
            self.document.add_paragraph(ai_insights['meeting_themes'])

        self.document.add_paragraph()

    def _add_footer_section(self, sprint_data: Dict[str, Any]) -> None:
        """Add report footer"""
        from datetime import datetime

        self.document.add_paragraph()
        footer = self.document.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} "
            f"by CSG Sprint Reporter"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
